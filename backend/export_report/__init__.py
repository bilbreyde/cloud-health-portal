import io
import json
import logging
from datetime import datetime, timezone

import azure.functions as func
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from shared import blob_client, cosmos_client
from shared.response_helpers import cors_options, cors_response
from shared.spend_insights_engine import (
    compute_anomalies,
    compute_commitment_utilization,
    compute_coverage_analysis,
    compute_correlations,
    compute_opportunities,
    last_n_months,
)
from shared.trend_engine import compute_mom_delta

_BLUE = RGBColor(0x17, 0x5E, 0x8C)
_GREY = RGBColor(0x60, 0x60, 0x60)
_TABLE_HDR_FILL = 'D5E8F0'
_FONT = 'Calibri'


def _prev_month(month: int, year: int):
    return (12, year - 1) if month == 1 else (month - 1, year)


def _cell_bg(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def _run(para, text, size=11, bold=False, color=None):
    r = para.add_run(text)
    r.font.name = _FONT
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    return r


def _para(doc, text='', size=11, bold=False, color=None, space_before=0, space_after=6, align=None):
    p = doc.add_paragraph()
    if text:
        _run(p, text, size=size, bold=bold, color=color)
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p


def _h1(doc, text):
    p = doc.add_paragraph()
    _run(p, text, size=14, bold=True, color=_BLUE)
    fmt = p.paragraph_format
    fmt.space_before = Pt(14)
    fmt.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '175E8C')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def _h2(doc, text):
    p = doc.add_paragraph()
    _run(p, text, size=12, bold=True, color=_BLUE)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    return p


def _tbl_header(table, headers):
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        cell.text = ''
        _run(cell.paragraphs[0], hdr, size=10, bold=True)
        _cell_bg(cell, _TABLE_HDR_FILL)


def _tbl_row(table, values):
    row = table.add_row()
    for i, val in enumerate(values):
        if i < len(row.cells):
            cell = row.cells[i]
            cell.text = ''
            _run(cell.paragraphs[0], str(val), size=10)
    return row


def _bullet(doc, text):
    p = doc.add_paragraph()
    _run(p, f'•  {text}', size=11)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    return p


def _multiline(doc, text, **kw):
    for para in (text or '').split('\n\n'):
        para = para.strip()
        if para:
            _para(doc, para, space_after=8, **kw)


def _fmt(n: float) -> str:
    return f'${n:,.2f}'


def _build_docx(
    customer_name: str,
    month: int,
    year: int,
    month_label: str,
    narrative: dict,
    curr_data: dict,
    prev_data: dict,
    top_movers_up: list,
    top_movers_down: list,
    service_summary: list,
    total_signal: float,
    exc_floor: float,
    net_addressable: float,
    realized_savings: float,
    remaining: float,
    exc_summary: dict | None,
    exc_records: list,
    cost_summary: dict | None,
    spend_insights: dict | None,
    prev_next_steps: list,
    ongoing_next_steps: list,
    planned_savings: list,
    project_updates: list,
    progress_narrative: str,
    joel_notes: str,
) -> bytes:
    doc = Document()

    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    # ── Title page ─────────────────────────────────────────────────────────────
    _para(doc, space_after=24)
    _para(doc, 'Cloud Cost Optimization Report', size=20, bold=True, color=_BLUE,
          space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, customer_name, size=16, bold=True,
          space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, month_label, size=14,
          space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, 'Prepared by: Dycom Technology Solutions', size=11, color=_GREY,
          space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, f'Report Generated: {datetime.now(timezone.utc).strftime("%B %d, %Y")}',
          size=11, color=_GREY, space_after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ── Section 1: Executive Summary ───────────────────────────────────────────
    _h1(doc, '1. Executive Summary')
    _multiline(doc, narrative.get('executive_summary', ''))

    _h2(doc, 'Savings Overview')
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    _tbl_header(tbl, ['Metric', 'Amount'])
    _tbl_row(tbl, ['CloudHealth Signal (Theoretical Maximum)', _fmt(total_signal)])
    _tbl_row(tbl, ['Exception Floor (Business-Critical Exclusions)', _fmt(exc_floor)])
    _tbl_row(tbl, ['Net Addressable Opportunity', _fmt(net_addressable)])
    if realized_savings > 0:
        _tbl_row(tbl, ['Realized Savings (Confirmed Executed)', _fmt(realized_savings)])
        _tbl_row(tbl, ['Remaining Opportunity', _fmt(remaining)])
    _para(doc)

    if top_movers_up or top_movers_down:
        _h2(doc, 'Month-over-Month Summary')
        tbl2 = doc.add_table(rows=1, cols=3)
        tbl2.style = 'Table Grid'
        _tbl_header(tbl2, ['Service', 'MoM Delta', 'Direction'])
        for m in top_movers_up[:5]:
            _tbl_row(tbl2, [m['serviceType'], f'+{_fmt(m["momDelta"])}', '▲ Up'])
        for m in top_movers_down[:5]:
            _tbl_row(tbl2, [m['serviceType'], _fmt(m['momDelta']), '▼ Down'])
        _para(doc)

    _h2(doc, 'Service Savings Summary')
    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = 'Table Grid'
    _tbl_header(tbl3, ['Service', 'Savings Signal', 'MoM Delta', 'Classification'])
    for row in sorted(service_summary, key=lambda x: -x['savingsTotal']):
        d = row['momDelta']
        _tbl_row(tbl3, [
            row['serviceType'],
            _fmt(row['savingsTotal']),
            ('+' if d >= 0 else '') + _fmt(d),
            row.get('classification') or '—',
        ])
    doc.add_page_break()

    # ── Section 2: Optimization Progress ──────────────────────────────────────
    _h1(doc, '2. Optimization Progress')
    _multiline(doc, narrative.get('optimization_narrative', ''))

    if progress_narrative:
        _h2(doc, '2.1 Progress Context')
        for line in progress_narrative.split('\n'):
            line = line.strip()
            if line:
                _para(doc, line, space_after=4)

    if prev_next_steps:
        _h2(doc, '2.2 Commitments from Previous Cycle')
        for step in prev_next_steps:
            _bullet(doc, step)

    if ongoing_next_steps:
        _h2(doc, '2.3 Ongoing Recurring Commitments')
        for step in ongoing_next_steps:
            _bullet(doc, step)

    if planned_savings:
        _h2(doc, '2.4 Upcoming Planned Savings Pipeline')
        for item in planned_savings:
            _bullet(doc, item)

    if project_updates:
        _h2(doc, '2.5 Active Projects & Migration Status')
        for item in project_updates:
            _bullet(doc, item)

    doc.add_page_break()

    # ── Section 3: AWS Total Spend Overview ───────────────────────────────────
    _h1(doc, '3. AWS Total Spend Overview')
    _para(doc,
          'Actual AWS billing for the reporting window, separate from the CloudHealth optimization '
          'signal reported above.', space_after=8)
    if narrative.get('aws_spend_overview'):
        _multiline(doc, narrative['aws_spend_overview'])

    if cost_summary and cost_summary.get('monthlyTotals'):
        _h2(doc, '3.1 Monthly Spend — Jul 2025 to Present')
        tbl_spend = doc.add_table(rows=1, cols=4)
        tbl_spend.style = 'Table Grid'
        _tbl_header(tbl_spend, ['Month', 'Direct Charges', 'Indirect Charges', 'Net Cost'])
        for m in cost_summary['monthlyTotals']:
            month_label = f"{m['month']} (partial, {m['completionRatio'] * 100:.0f}% complete)" \
                if m.get('isPartial') else m['month']
            _tbl_row(tbl_spend, [
                month_label, _fmt(m['directCharges']), _fmt(m['indirectCharges']), _fmt(m['netCost']),
            ])
        if any(m.get('isPartial') for m in cost_summary['monthlyTotals']):
            _para(doc, 'The partial month above shows actual spend to date, not a projection; '
                       'MoM comparisons elsewhere in this report use the projected full-month figure.',
                  size=9, color=_GREY)
        _para(doc)

        _h2(doc, '3.2 Top 10 Services by Cost')
        tbl_top = doc.add_table(rows=1, cols=4)
        tbl_top.style = 'Table Grid'
        _tbl_header(tbl_top, ['Service', 'Current Month', 'Prior Month', 'MoM Delta'])
        for s in cost_summary['topServices']:
            d = s['momDelta']
            _tbl_row(tbl_top, [
                s['service'], _fmt(s['currentMonth']), _fmt(s['previousMonth']),
                ('+' if d >= 0 else '') + _fmt(d),
            ])
        _para(doc)

        _h2(doc, '3.3 Savings Plan Coverage')
        cov = cost_summary['savingsPlanCoverage']
        tbl_cov = doc.add_table(rows=1, cols=2)
        tbl_cov.style = 'Table Grid'
        _tbl_header(tbl_cov, ['Metric', 'Value'])
        _tbl_row(tbl_cov, ['Savings Plan Covered Spend', _fmt(cov['covered'])])
        _tbl_row(tbl_cov, ['On-Demand Spend', _fmt(cov['onDemand'])])
        _tbl_row(tbl_cov, ['Coverage %', f"{cov['coveragePct']}%"])
        if cost_summary.get('projectedCurrentMonth'):
            _tbl_row(tbl_cov, ['Projected Current Month Spend', _fmt(cost_summary['projectedCurrentMonth'])])
    else:
        _para(doc, 'No AWS Cost History data imported for this customer yet.')
    doc.add_page_break()

    # ── Section 4: AWS Spend Analysis & Optimization Opportunities ────────────
    _h1(doc, '4. AWS Spend Analysis & Optimization Opportunities')
    if spend_insights:
        if spend_insights.get('narrative'):
            _multiline(doc, spend_insights['narrative'])

        _h2(doc, '4.1 Anomalies Detected')
        anomalies = spend_insights.get('anomalies') or []
        if anomalies:
            tbl_anom = doc.add_table(rows=1, cols=4)
            tbl_anom.style = 'Table Grid'
            _tbl_header(tbl_anom, ['Service', 'Current Amount', 'Flag Type', 'Explanation'])
            for a in anomalies[:10]:
                flag = a.get('flagType') or a['type'].replace('_', ' ').title()
                _tbl_row(tbl_anom, [a['service'], _fmt(a['currentAmount']), flag, a['explanation']])
        else:
            _para(doc, 'No statistically significant anomalies detected this period.')
        _para(doc)

        cov_sp = spend_insights.get('coverageAnalysis') or {}
        commitment_util = spend_insights.get('commitmentUtilization') or {}

        # Notable one-time charges, separate from recurring/anomaly analysis — sourced
        # from the commitment breakdown when available, else from one-time-pattern anomalies.
        _h2(doc, '4.2 Notable One-Time Charges')
        one_time_rows = commitment_util.get('excludedServices') or [
            {'service': a['service'], 'amount': a['currentAmount'], 'reason': a.get('flagType', 'One-Time')}
            for a in anomalies if a.get('pattern') == 'one_time'
        ]
        if one_time_rows:
            tbl_ot = doc.add_table(rows=1, cols=3)
            tbl_ot.style = 'Table Grid'
            _tbl_header(tbl_ot, ['Service', 'Amount', 'Reason'])
            for r in one_time_rows[:10]:
                _tbl_row(tbl_ot, [r['service'], _fmt(r['amount']), r['reason']])
            _para(doc, 'Excluded from EDP/commitment utilization and not extrapolated for projection.',
                  size=9, color=_GREY)
        else:
            _para(doc, 'No notable one-time charges this period.')
        _para(doc)

        if commitment_util:
            _h2(doc, '4.3 Commitment (EDP) Utilization')
            tbl_cu = doc.add_table(rows=1, cols=2)
            tbl_cu.style = 'Table Grid'
            _tbl_header(tbl_cu, ['Metric', 'Value'])
            _tbl_row(tbl_cu, ['Commitment Type', commitment_util.get('commitmentType') or '—'])
            _tbl_row(tbl_cu, ['Monthly Obligation', _fmt(commitment_util.get('monthlyObligation', 0))])
            _tbl_row(tbl_cu, ['Recurring Spend (vs. obligation)', _fmt(commitment_util.get('recurringSpend', 0))])
            _tbl_row(tbl_cu, ['Utilization %', f"{commitment_util.get('utilizationPct', 0)}%"])
            _tbl_row(tbl_cu, ['One-Time Charges (excluded)', _fmt(commitment_util.get('oneTimeCharges', 0))])
            _tbl_row(tbl_cu, ['Credits Applied', f"-{_fmt(commitment_util.get('credits', 0))}"])
            _tbl_row(tbl_cu, ['Net Billed', _fmt(commitment_util.get('netBilled', 0))])
            _tbl_row(tbl_cu, ['Status', 'On Track' if commitment_util.get('onTrack') else 'Off Track'])
            over_under = commitment_util.get('overUnderAmount', 0)
            _tbl_row(tbl_cu, [
                'Recurring Over / Under Obligation',
                (('+' if over_under >= 0 else '') + _fmt(over_under)),
            ])
            _tbl_row(tbl_cu, ['Trailing 3-Month Recurring Average', _fmt(commitment_util.get('trailing3MoAvg', 0))])
            months_remaining = commitment_util.get('monthsRemaining')
            if months_remaining is not None:
                _tbl_row(tbl_cu, ['Months Remaining on Commitment', str(months_remaining)])
            _para(doc)
            if commitment_util.get('underUtilizationRisk'):
                _para(doc, 'Risk: trailing 3-month recurring average is below 85% of the monthly '
                           'obligation — this commitment is at risk of under-utilization.')
            if commitment_util.get('expiryWarning'):
                _para(doc, 'Renewal decision needed — commitment expires within 6 months; '
                           'a current pricing vs. market rate analysis is recommended.')
        else:
            _h2(doc, '4.3 Savings Plan Coverage Recommendation')
            if cov_sp:
                tbl_sp = doc.add_table(rows=1, cols=2)
                tbl_sp.style = 'Table Grid'
                _tbl_header(tbl_sp, ['Metric', 'Value'])
                _tbl_row(tbl_sp, ['Current Coverage', f"{cov_sp.get('currentPct', 0)}%"])
                _tbl_row(tbl_sp, ['Target Coverage', f"{cov_sp.get('targetPct', 0)}%"])
                _tbl_row(tbl_sp, ['Gap (On-Demand → Committed)', _fmt(cov_sp.get('gapAmount', 0))])
                _tbl_row(tbl_sp, ['Estimated Additional Savings', _fmt(cov_sp.get('estimatedSavings', 0)) + '/mo'])
                rec = cov_sp.get('recommendation') or {}
                if rec.get('term'):
                    _tbl_row(tbl_sp, ['Recommended Term', rec['term']])
                _para(doc)
                if rec.get('rationale'):
                    _para(doc, rec['rationale'])
        _para(doc)

        _h2(doc, '4.4 Optimization Opportunities')
        opportunities = spend_insights.get('opportunities') or []
        if opportunities:
            tbl_opp = doc.add_table(rows=1, cols=5)
            tbl_opp.style = 'Table Grid'
            _tbl_header(tbl_opp, ['Priority', 'Service', 'Category', 'Est. Savings/mo', 'Recommended Action'])
            for o in opportunities[:10]:
                savings = _fmt(o['estimatedSavings']) if o['estimatedSavings'] > 0 else 'N/A — risk mitigation'
                _tbl_row(tbl_opp, [o['priority'], o['service'], o['category'], savings, o['action']])
        else:
            _para(doc, 'No additional optimization opportunities identified beyond CloudHealth signal.')
    else:
        _para(doc, 'AI spend insights have not been generated for this customer yet.')
    doc.add_page_break()

    # ── Section 5: Service Level Analysis ─────────────────────────────────────
    _h1(doc, '5. Service Level Analysis')
    _para(doc, 'Detailed breakdown of CloudHealth savings signal by AWS service for the reporting period.',
          space_after=8)
    tbl_svc = doc.add_table(rows=1, cols=5)
    tbl_svc.style = 'Table Grid'
    _tbl_header(tbl_svc, ['Service', 'Current Signal', 'Prior Month', 'MoM Delta', 'Direction'])
    for row in sorted(service_summary, key=lambda x: -x['savingsTotal']):
        prev_val = prev_data.get(row['serviceType'], 0.0)
        d = row['momDelta']
        _tbl_row(tbl_svc, [
            row['serviceType'],
            _fmt(row['savingsTotal']),
            _fmt(prev_val),
            ('+' if d >= 0 else '') + _fmt(d),
            row.get('direction', '—'),
        ])
    doc.add_page_break()

    # ── Section 6: Top Movers Analysis ────────────────────────────────────────
    _h1(doc, '6. Top Movers Analysis')
    _multiline(doc, narrative.get('top_movers_analysis', ''))

    _h2(doc, '6.1 Spending Increases')
    if top_movers_up:
        tbl_up = doc.add_table(rows=1, cols=3)
        tbl_up.style = 'Table Grid'
        _tbl_header(tbl_up, ['Service', 'MoM Increase', 'Classification'])
        for m in top_movers_up:
            cls = next((s.get('classification', '') for s in service_summary
                        if s['serviceType'] == m['serviceType']), '')
            _tbl_row(tbl_up, [m['serviceType'], f'+{_fmt(m["momDelta"])}', cls or '—'])
    else:
        _para(doc, 'No significant spending increases this period.')
    _para(doc)

    _h2(doc, '6.2 Spending Decreases & Savings Realized')
    if top_movers_down:
        tbl_dn = doc.add_table(rows=1, cols=3)
        tbl_dn.style = 'Table Grid'
        _tbl_header(tbl_dn, ['Service', 'MoM Decrease', 'Classification'])
        for m in top_movers_down:
            cls = next((s.get('classification', '') for s in service_summary
                        if s['serviceType'] == m['serviceType']), '')
            _tbl_row(tbl_dn, [m['serviceType'], _fmt(m['momDelta']), cls or '—'])
    else:
        _para(doc, 'No significant spending decreases this period.')
    doc.add_page_break()

    # ── Section 7: EC2 Deep Dive ───────────────────────────────────────────────
    ec2_signal = curr_data.get('EC2', 0.0)
    if ec2_signal > 0:
        _h1(doc, '7. EC2 Deep Dive')
        ec2_row = next((s for s in service_summary if s['serviceType'] == 'EC2'), None)
        if ec2_row:
            tbl_ec2 = doc.add_table(rows=1, cols=2)
            tbl_ec2.style = 'Table Grid'
            _tbl_header(tbl_ec2, ['Metric', 'Value'])
            _tbl_row(tbl_ec2, ['Current Month Signal', _fmt(ec2_signal)])
            _tbl_row(tbl_ec2, ['Prior Month Signal', _fmt(prev_data.get('EC2', 0.0))])
            d = ec2_row['momDelta']
            _tbl_row(tbl_ec2, ['Month-over-Month Delta', ('+' if d >= 0 else '') + _fmt(d)])
            _tbl_row(tbl_ec2, ['Trend Direction', ec2_row.get('direction', '—')])
        _para(doc,
              'EC2 is typically the primary driver of cloud compute spend. '
              'See the Exception Register (Section 9) for instances excluded from optimization scope.',
              space_before=8)
        doc.add_page_break()

    # ── Section 8: Savings Breakdown ──────────────────────────────────────────
    _h1(doc, '8. Savings Breakdown')
    _multiline(doc, narrative.get('exception_delta', ''))

    tbl_flow = doc.add_table(rows=1, cols=2)
    tbl_flow.style = 'Table Grid'
    _tbl_header(tbl_flow, ['Category', 'Amount'])
    _tbl_row(tbl_flow, ['Gross CloudHealth Signal', _fmt(total_signal)])
    _tbl_row(tbl_flow, ['Less: Exception Floor', f'({_fmt(exc_floor)})'])
    _tbl_row(tbl_flow, ['Net Addressable Opportunity', _fmt(net_addressable)])
    if realized_savings > 0:
        _tbl_row(tbl_flow, ['Less: Realized Savings', f'({_fmt(realized_savings)})'])
        _tbl_row(tbl_flow, ['Remaining Opportunity', _fmt(remaining)])
    doc.add_page_break()

    # ── Section 9: Exceptions ─────────────────────────────────────────────────
    _h1(doc, '9. Exceptions Register')

    _h2(doc, '9.1 Exception Fleet Summary')
    if exc_summary and exc_summary.get('totalCount', 0) > 0:
        _para(doc,
              f"Total exception servers: {exc_summary['totalCount']} | "
              f"Total monthly cost: {_fmt(exc_summary['totalMonthlyCost'])}",
              space_after=6)
        tbl_cat = doc.add_table(rows=1, cols=3)
        tbl_cat.style = 'Table Grid'
        _tbl_header(tbl_cat, ['Category', 'Server Count', 'Monthly Cost'])
        for cat in (exc_summary.get('byCategory') or []):
            _tbl_row(tbl_cat, [cat['category'], str(cat['count']), _fmt(cat['monthlyCost'])])
        _para(doc)
        if exc_summary.get('byLifecycle'):
            _h2(doc, 'Lifecycle Breakdown')
            tbl_lc = doc.add_table(rows=1, cols=3)
            tbl_lc.style = 'Table Grid'
            _tbl_header(tbl_lc, ['Lifecycle', 'Server Count', 'Monthly Cost'])
            for lc in exc_summary['byLifecycle']:
                _tbl_row(tbl_lc, [lc['lifecycle'], str(lc['count']), _fmt(lc['monthlyCost'])])
    else:
        _para(doc, 'No exception servers recorded for this customer.')
    _para(doc)

    _h2(doc, '9.2 Exception Register')
    if exc_records:
        tbl_exc = doc.add_table(rows=1, cols=6)
        tbl_exc.style = 'Table Grid'
        _tbl_header(tbl_exc, ['Instance Name', 'Account', 'Category', 'Lifecycle', 'Monthly Cost', 'Notes'])
        for rec in exc_records[:50]:
            _tbl_row(tbl_exc, [
                rec.instanceName or rec.instanceId,
                rec.accountName or '—',
                rec.exceptionCategory or '—',
                rec.lifecycle or '—',
                _fmt(rec.projectedCostPerMonth),
                (rec.notes or '')[:80],
            ])
    else:
        _para(doc, 'No exception records on file.')
    _para(doc)

    _h2(doc, '9.3 Exception & Signal Delta Analysis')
    exc_text = narrative.get('exception_delta', '')
    if not exc_text:
        exc_text = (
            f'The CloudHealth signal of {_fmt(total_signal)} represents the theoretical maximum savings '
            f'opportunity. The exception floor of {_fmt(exc_floor)} represents business-critical servers '
            f'excluded from optimization scope. The net addressable opportunity is {_fmt(net_addressable)}.'
        )
    _multiline(doc, exc_text)
    doc.add_page_break()

    # ── Section 10: Risks & Constraints ────────────────────────────────────────
    _h1(doc, '10. Risks & Constraints')
    _multiline(doc, narrative.get('risks_and_next_steps', ''))
    doc.add_page_break()

    # ── Section 11: Next Steps ────────────────────────────────────────────────
    _h1(doc, '11. Next Steps')

    _h2(doc, 'Before Next Meeting')
    if prev_next_steps:
        for step in prev_next_steps:
            _bullet(doc, step)
    else:
        _para(doc, 'No specific one-time actions defined for next cycle.')
    _para(doc)

    _h2(doc, 'Ongoing Commitments')
    if ongoing_next_steps:
        for step in ongoing_next_steps:
            _bullet(doc, step)
    else:
        _para(doc, 'No recurring commitments defined.')
    doc.add_page_break()

    # ── Section 12: Appendix ──────────────────────────────────────────────────
    _h1(doc, '12. Appendix')

    if joel_notes:
        _h2(doc, 'A. Engagement Manager Notes')
        _para(doc, joel_notes, space_after=8)

    _h2(doc, 'B. Report Methodology')
    _para(doc,
          'CloudHealth signal represents savings opportunities identified through automated analysis of AWS '
          'resource utilization, right-sizing recommendations, and cost optimization rules. '
          'The exception floor represents servers excluded from optimization scope due to business-critical '
          'designation, licensing constraints, or active vendor negotiations. '
          'Realized savings represent confirmed, executed optimization actions only, '
          'and are not inferred from signal movement alone.')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('export_report triggered')
    if req.method == 'OPTIONS':
        return cors_options()
    try:
        body = req.get_json()
    except ValueError:
        return cors_response({'error': 'Request body must be valid JSON'}, 400)

    customer_id = (body.get('customerId') or '').strip()
    month = body.get('month')
    year = body.get('year')

    if not customer_id:
        return cors_response({'error': 'customerId is required'}, 400)
    if not isinstance(month, int) or not isinstance(year, int):
        return cors_response({'error': 'month and year must be integers'}, 400)
    if not 1 <= month <= 12:
        return cors_response({'error': 'month must be 1-12'}, 400)

    customer = cosmos_client.get_customer(customer_id)
    if customer is None:
        return cors_response({'error': f'Customer {customer_id!r} not found'}, 404)

    all_trends = cosmos_client.list_trends(customer_id)
    curr_trends = [t for t in all_trends if t.year == year and t.month == month]
    if not curr_trends:
        return cors_response(
            {'error': f'No trend data found for {month}/{year}. Please upload CSVs for this period first.'},
            404,
        )

    prev_month_n, prev_year = _prev_month(month, year)
    prev_trends = [t for t in all_trends if t.year == prev_year and t.month == prev_month_n]

    curr_data = {t.serviceType: t.savingsTotal for t in curr_trends}
    prev_data = {t.serviceType: t.savingsTotal for t in prev_trends}

    # ── Report data ────────────────────────────────────────────────────────────
    narrative: dict = {
        'executive_summary': '', 'optimization_narrative': '',
        'top_movers_analysis': '', 'risks_and_next_steps': '', 'exception_delta': '',
        'aws_spend_overview': '',
    }
    joel_notes: str = ''
    realized_savings: float = 0.0
    prev_next_steps: list = []
    ongoing_next_steps: list = []
    planned_savings: list = []
    project_updates: list = []
    progress_narrative: str = ''

    try:
        all_reports = cosmos_client.list_reports(customer_id)
        generated = next(
            (r for r in all_reports if r.source == 'generated' and r.year == year and r.month == month),
            None,
        ) or next((r for r in all_reports if r.source == 'generated'), None)

        if generated:
            if generated.narrativeDraft:
                narrative = json.loads(generated.narrativeDraft)
            joel_notes = generated.joelNotes or ''

        imported_curr = next(
            (r for r in all_reports
             if r.source == 'manual_import' and r.year == year and r.month == month), None
        )
        imported_prev = next(
            (r for r in all_reports
             if r.source == 'manual_import' and r.year == prev_year and r.month == prev_month_n), None
        )

        if imported_curr and imported_curr.extractedData:
            realized_savings = float(imported_curr.extractedData.get('realizedSavings', 0.0))

        if imported_prev and imported_prev.extractedData:
            ed = imported_prev.extractedData
            prev_next_steps = ed.get('nextSteps', []) or []
            ongoing_next_steps = ed.get('ongoingNextSteps', []) or []
            planned_savings = ed.get('plannedSavings', []) or []
            project_updates = ed.get('projectUpdates', []) or []
            progress_narrative = ed.get('progressNarrative', '') or ''
            if not prev_data:
                prev_data = ed.get('monthlySavings', {})
    except Exception as exc:
        logging.warning('Report data lookup failed (non-fatal): %s', exc)

    # ── MoM deltas ────────────────────────────────────────────────────────────
    top_movers_up: list = []
    top_movers_down: list = []
    service_summary: list = []

    for svc, curr_total in curr_data.items():
        prev_total = prev_data.get(svc, 0.0)
        delta, direction = compute_mom_delta(curr_total, prev_total)
        classification = ''
        if abs(delta) > 500:
            classification = 'Persistent Issue' if svc in prev_data else 'New Insight'
        service_summary.append({
            'serviceType': svc, 'savingsTotal': curr_total,
            'momDelta': delta, 'direction': direction, 'classification': classification,
        })
        if delta > 500:
            top_movers_up.append({'serviceType': svc, 'momDelta': delta})
        elif delta < -500:
            top_movers_down.append({'serviceType': svc, 'momDelta': delta})

    top_movers_up.sort(key=lambda x: -x['momDelta'])
    top_movers_down.sort(key=lambda x: x['momDelta'])

    # ── Exception data ─────────────────────────────────────────────────────────
    exc_summary = None
    exc_records = []
    try:
        exc_summary = cosmos_client.exceptions_summary(customer_id)
        exc_records = cosmos_client.list_exceptions(customer_id)
    except Exception as exc:
        logging.warning('Exception data fetch failed (non-fatal): %s', exc)

    # ── Cost history data ───────────────────────────────────────────────────────
    cost_summary = None
    try:
        all_cost_records = cosmos_client.get_cost_history(customer_id, '0000-00', '9999-99')
        cost_months = sorted({r.month for r in all_cost_records})
        if cost_months:
            cost_summary = cosmos_client.get_cost_history_summary(customer_id, cost_months)
    except Exception as exc:
        logging.warning('Cost history fetch failed (non-fatal): %s', exc)

    # ── AI spend insights ──────────────────────────────────────────────────────
    # Reuse the cached analysis from the spend-insights page if it's there; otherwise
    # recompute the structured tables fresh (cheap — no AI call) rather than skip the
    # section, but don't trigger a fresh AI narrative call during export.
    spend_insights = None
    try:
        insights_month = f'{year:04d}-{month:02d}'
        cached_insights = cosmos_client.get_report(f'insights-{customer_id}-{insights_month}', customer_id)
        if cached_insights and cached_insights.source == 'spend_insights' and cached_insights.extractedData:
            spend_insights = cached_insights.extractedData
        elif cost_summary and cost_months:
            window_months = last_n_months(insights_month, 6)
            months_in_window = [m for m in window_months if m in cost_months]
            if months_in_window:
                window_summary = cosmos_client.get_cost_history_summary(customer_id, months_in_window)
                si_by_service = window_summary['byService']
                si_is_partial = window_summary['isPartial']
                si_completion_ratio = window_summary['completionRatio']
                all_trends_for_insights = cosmos_client.list_trends(customer_id)
                trend_dicts = [
                    {'serviceType': t.serviceType, 'month': t.month, 'year': t.year, 'savingsTotal': t.savingsTotal}
                    for t in all_trends_for_insights
                    if f'{t.year:04d}-{t.month:02d}' in months_in_window
                ]
                si_anomalies = compute_anomalies(si_by_service, months_in_window, si_is_partial, si_completion_ratio)
                si_correlations = compute_correlations(
                    si_by_service, trend_dicts, months_in_window, si_is_partial, si_completion_ratio)

                si_commitment = (customer.settings or {}).get('commitment') or {}
                si_commitment_type = si_commitment.get('commitmentType')
                si_has_commitment = bool(si_commitment_type) and si_commitment_type != 'None'
                si_coverage = None
                si_commitment_utilization = None
                if si_has_commitment:
                    si_commitment_utilization = compute_commitment_utilization(
                        si_commitment, window_summary['monthlyTotals'], months_in_window, si_by_service,
                        si_is_partial, si_completion_ratio,
                    )
                else:
                    si_coverage = compute_coverage_analysis(
                        window_summary['savingsPlanCoverage'], si_by_service, months_in_window,
                        si_is_partial, si_completion_ratio,
                    )

                si_monthly_obligation = (
                    si_commitment_utilization['monthlyObligation'] if si_commitment_utilization else None
                )
                si_opportunities = compute_opportunities(
                    si_by_service, si_coverage, months_in_window[-1],
                    months_in_window=months_in_window,
                    monthly_obligation=si_monthly_obligation,
                    skip_savings_plan_opportunity=si_has_commitment,
                    is_partial=si_is_partial, completion_ratio=si_completion_ratio,
                )
                report_for_month = locals().get('generated')
                saved_narrative = (
                    (report_for_month.extractedData or {}).get('spendInsightsNarrative', '')
                    if report_for_month else ''
                )
                spend_insights = {
                    'anomalies': si_anomalies,
                    'coverageAnalysis': si_coverage,
                    'commitmentUtilization': si_commitment_utilization,
                    'correlations': si_correlations,
                    'opportunities': si_opportunities,
                    'narrative': saved_narrative,
                }
    except Exception as exc:
        logging.warning('Spend insights fetch failed (non-fatal): %s', exc)

    exc_floor = exc_summary['totalMonthlyCost'] if exc_summary else 0.0
    total_signal = sum(curr_data.values())
    net_addressable = max(0.0, total_signal - exc_floor)
    remaining = max(0.0, net_addressable - realized_savings)
    month_label = datetime(year, month, 1).strftime('%B %Y')

    # ── Build document ─────────────────────────────────────────────────────────
    try:
        docx_bytes = _build_docx(
            customer_name=customer.name,
            month=month,
            year=year,
            month_label=month_label,
            narrative=narrative,
            curr_data=curr_data,
            prev_data=prev_data,
            top_movers_up=top_movers_up,
            top_movers_down=top_movers_down,
            service_summary=service_summary,
            total_signal=total_signal,
            exc_floor=exc_floor,
            net_addressable=net_addressable,
            realized_savings=realized_savings,
            remaining=remaining,
            exc_summary=exc_summary,
            exc_records=exc_records,
            cost_summary=cost_summary,
            spend_insights=spend_insights,
            prev_next_steps=prev_next_steps,
            ongoing_next_steps=ongoing_next_steps,
            planned_savings=planned_savings,
            project_updates=project_updates,
            progress_narrative=progress_narrative,
            joel_notes=joel_notes,
        )
    except Exception as exc:
        logging.exception('docx build failed')
        return cors_response({'error': f'Document generation failed: {exc}'}, 500)

    # ── Save to blob (best-effort) ─────────────────────────────────────────────
    safe_name = f'{customer.slug}_{year}_{month:02d}_report.docx'
    try:
        blob_client.upload_docx(customer_id, month, year, docx_bytes, safe_name)
    except Exception as exc:
        logging.warning('Blob save failed (non-fatal): %s', exc)

    return func.HttpResponse(
        body=docx_bytes,
        status_code=200,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={
            'Access-Control-Allow-Origin': '*',
            'Content-Disposition': f'attachment; filename="{safe_name}"',
            'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'Access-Control-Expose-Headers': 'Content-Disposition',
        },
    )
