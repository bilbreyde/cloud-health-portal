import json
import logging
import re
import uuid
from datetime import datetime, timezone
from io import BytesIO

import azure.functions as func

from shared import blob_client, cosmos_client
from shared.models import Report

_DOLLAR_RE = re.compile(r'\$?\s*([\d,]+(?:\.\d{1,2})?)')
_SECTION_RE = re.compile(r'^(?:Section\s+)?(\d+\.\d+)\b', re.IGNORECASE)
_MAJOR_RE   = re.compile(r'^(?:Section\s+)?(\d+)\.?\s+[A-Z]')
_BULLET_RE  = re.compile(r'^[-•*●◦▪▸\d]+[.)]\s*')
_SERVICE_TYPES = {'EC2', 'EBS', 'RDS', 'S3', 'ElastiCache', 'Redshift',
                  'OpenSearch', 'DynamoDB', 'Consolidated'}


def _extract_dollar(text: str) -> float:
    m = _DOLLAR_RE.search(text)
    if not m:
        return 0.0
    try:
        return float(m.group(1).replace(',', ''))
    except (ValueError, IndexError):
        return 0.0


def _match_service(text: str) -> str | None:
    for svc in _SERVICE_TYPES:
        if svc.lower() in text.lower():
            return svc
    return None


def _is_new_major_section(text: str) -> bool:
    """Returns True if the paragraph looks like a top-level section header."""
    return bool(_MAJOR_RE.match(text)) and not bool(_SECTION_RE.match(text))


def _strip_bullet(text: str) -> str:
    return _BULLET_RE.sub('', text).strip()


def _parse_docx(file_bytes: bytes) -> dict:
    from docx import Document

    extracted: dict = {
        'monthlySavings': {},
        'topMoversUp': [],
        'topMoversDown': [],
        'realizedSavings': 0.0,
        'exceptionFloor': 0.0,
        'nextSteps': [],           # "Before Next Meeting" one-time items
        'ongoingNextSteps': [],    # "Ongoing" recurring commitments
        'plannedSavings': [],      # Upcoming pipeline items (section 2.3)
        'projectUpdates': [],      # Migration / project status items
        'progressNarrative': '',   # Section 2 prose paragraphs
    }

    doc = Document(BytesIO(file_bytes))

    # ── Tables: savings signal table ───────────────────────────────────────────
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [c.text.strip() for c in table.rows[0].cells]
        header_text = ' '.join(headers).lower()
        if not ('service' in header_text or 'saving' in header_text or
                any(s.lower() in header_text for s in _SERVICE_TYPES)):
            continue
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if not cells or not cells[0]:
                continue
            svc = _match_service(cells[0])
            if svc:
                for cell_val in reversed(cells[1:]):
                    amt = _extract_dollar(cell_val)
                    if amt > 0:
                        extracted['monthlySavings'][svc] = amt
                        break

    # ── Paragraphs: state machine ──────────────────────────────────────────────
    # Modes
    IN_NONE          = 'none'
    IN_SECTION2      = 'section2'       # generic section 2 prose
    IN_PLANNED       = 'planned'        # 2.3 or keyword "Upcoming Planned Savings"
    IN_PROJECTS      = 'projects'       # 2.4 / migration / project updates
    IN_MOVERS_UP     = 'movers_up'
    IN_MOVERS_DOWN   = 'movers_down'
    IN_STEPS_BEFORE  = 'steps_before'   # Section 9, "Before Next Meeting"
    IN_STEPS_ONGOING = 'steps_ongoing'  # Section 9, "Ongoing"

    mode = IN_NONE
    progress_parts: list[str] = []

    def set_mode(new_mode: str) -> str:
        return new_mode

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        lower = text.lower()

        # ── Numbered subsection header (e.g. "2.3 Upcoming Planned Savings") ──
        sec_m = _SECTION_RE.match(text)
        if sec_m:
            sec  = sec_m.group(1)
            major = sec.split('.')[0]
            rest  = text[sec_m.end():].strip().lower()

            if major == '2':
                if re.search(r'planned\s+sav|upcoming|pipeline', rest):
                    mode = IN_PLANNED
                elif re.search(r'migrat|project|status|update|in.progress', rest):
                    mode = IN_PROJECTS
                else:
                    mode = IN_SECTION2
            elif sec == '4.1' or re.search(r'mover.*(up|increas)', rest):
                mode = IN_MOVERS_UP
            elif sec == '4.2' or re.search(r'mover.*(down|decreas|reduc)', rest):
                mode = IN_MOVERS_DOWN
            elif major == '9' or re.search(r'next\s+step|recommended', rest):
                mode = IN_STEPS_BEFORE
            else:
                mode = IN_NONE
            continue

        # ── Major section header (e.g. "9. Next Steps") ───────────────────────
        if _is_new_major_section(text):
            if re.search(r'next\s+step|action', lower):
                mode = IN_STEPS_BEFORE
            elif re.search(r'^2\b', text):
                mode = IN_SECTION2
            else:
                mode = IN_NONE
            continue

        # ── Keyword-based mode transitions (no section numbers) ───────────────
        if mode not in (IN_STEPS_BEFORE, IN_STEPS_ONGOING):
            if re.search(r'upcoming\s+planned\s+sav|planned\s+sav.*pipeline', lower) and len(text) < 80:
                mode = IN_PLANNED
                continue
            if re.search(r'top\s+movers?\s+(up|increas)', lower) and len(text) < 80:
                mode = IN_MOVERS_UP
                continue
            if re.search(r'top\s+movers?\s+(down|decreas|reduc)', lower) and len(text) < 80:
                mode = IN_MOVERS_DOWN
                continue
            if re.search(r'next\s+steps?\s*$|recommended\s+actions?\s*$', lower) and len(text) < 60:
                mode = IN_STEPS_BEFORE
                continue

        # ── Within Next Steps — detect Ongoing vs Before subsection ───────────
        if mode in (IN_STEPS_BEFORE, IN_STEPS_ONGOING):
            if re.search(r'\bongoin\b', lower) and len(text) < 50:
                mode = IN_STEPS_ONGOING
                continue
            if re.search(r'before\s+(next\s+)?meet', lower) and len(text) < 80:
                mode = IN_STEPS_BEFORE
                continue
            if _is_new_major_section(text):
                mode = IN_NONE
                continue
            item = _strip_bullet(text)
            if item and len(item) > 8:
                if mode == IN_STEPS_ONGOING:
                    if item not in extracted['ongoingNextSteps']:
                        extracted['ongoingNextSteps'].append(item)
                else:
                    if item not in extracted['nextSteps']:
                        extracted['nextSteps'].append(item)
            continue

        # ── Planned savings ───────────────────────────────────────────────────
        if mode == IN_PLANNED:
            if _is_new_major_section(text) or sec_m:
                pass  # will be caught above on next iter
            else:
                item = _strip_bullet(text)
                if item and len(item) > 5 and item not in extracted['plannedSavings']:
                    extracted['plannedSavings'].append(item)
            continue

        # ── Project / migration updates ────────────────────────────────────────
        if mode == IN_PROJECTS:
            if _is_new_major_section(text):
                pass
            else:
                item = _strip_bullet(text)
                if item and len(item) > 5 and item not in extracted['projectUpdates']:
                    extracted['projectUpdates'].append(item)
            continue

        # ── Top movers ────────────────────────────────────────────────────────
        if mode == IN_MOVERS_UP:
            svc = _match_service(text)
            amt = _extract_dollar(text)
            if svc and amt > 0 and not any(m['serviceType'] == svc for m in extracted['topMoversUp']):
                extracted['topMoversUp'].append({'serviceType': svc, 'amount': amt})
            continue

        if mode == IN_MOVERS_DOWN:
            svc = _match_service(text)
            amt = _extract_dollar(text)
            if svc and amt > 0 and not any(m['serviceType'] == svc for m in extracted['topMoversDown']):
                extracted['topMoversDown'].append({'serviceType': svc, 'amount': amt})
            continue

        # ── Section 2 prose → progressNarrative ───────────────────────────────
        if mode == IN_SECTION2:
            if not _BULLET_RE.match(text) and len(text) > 40:
                progress_parts.append(text)
            # Also look for migration/project mentions in free prose
            if re.search(
                r'\b(fsx|migrat|on\s+hold|vendor\s+meet|domain\s+controller|'
                r'pending\s+terminat|dxc|in\s+progress|scheduled)\b', lower
            ):
                item = text.strip()
                if len(item) > 15 and item not in extracted['projectUpdates']:
                    extracted['projectUpdates'].append(item)
            continue

        # ── Inline keyword extractions (anywhere in doc) ──────────────────────
        if re.search(r'realized\s+saving', lower):
            amt = _extract_dollar(text)
            if amt > 0:
                extracted['realizedSavings'] = amt

        if re.search(r'exception\s+floor|exception\s+total', lower):
            amt = _extract_dollar(text)
            if amt > 0:
                extracted['exceptionFloor'] = amt

    # Cap lists
    extracted['nextSteps']        = extracted['nextSteps'][:20]
    extracted['ongoingNextSteps'] = extracted['ongoingNextSteps'][:10]
    extracted['topMoversUp']      = extracted['topMoversUp'][:5]
    extracted['topMoversDown']    = extracted['topMoversDown'][:5]
    extracted['plannedSavings']   = extracted['plannedSavings'][:15]
    extracted['projectUpdates']   = extracted['projectUpdates'][:10]
    extracted['progressNarrative'] = '\n'.join(progress_parts[:6])

    return extracted


def _json(body, status: int = 200) -> func.HttpResponse:
    return func.HttpResponse(json.dumps(body), status_code=status, mimetype='application/json')


def main(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('import_report triggered')
    customer_id = (req.route_params.get('customerId') or '').strip()
    if not customer_id:
        return _json({'error': 'customerId is required'}, 400)

    try:
        return _handle(req, customer_id)
    except Exception as exc:
        logging.exception('import_report unhandled error')
        return _json({'error': str(exc)}, 500)


def _handle(req: func.HttpRequest, customer_id: str) -> func.HttpResponse:
    customer = cosmos_client.get_customer(customer_id)
    if customer is None:
        return _json({'error': f'Customer {customer_id!r} not found'}, 404)

    def field(name: str) -> str:
        return (req.params.get(name) or req.form.get(name, '')).strip()

    month_str = field('month')
    year_str  = field('year')
    field('reportDate')

    if not month_str or not year_str:
        return _json({'error': 'month and year are required'}, 400)
    try:
        month = int(month_str)
        year  = int(year_str)
    except ValueError:
        return _json({'error': 'month and year must be integers'}, 400)
    if not 1 <= month <= 12:
        return _json({'error': 'month must be 1–12'}, 400)

    uploaded = req.files.get('file')
    if uploaded is None:
        return _json({'error': 'multipart field "file" is required'}, 400)

    filename = uploaded.filename or 'report.docx'
    if not filename.lower().endswith('.docx'):
        return _json({'error': 'Only .docx files are supported'}, 400)

    file_bytes = uploaded.read()

    try:
        extracted_data = _parse_docx(file_bytes)
    except Exception as exc:
        logging.warning('docx parse error (storing empty extractedData): %s', exc)
        extracted_data = {
            'monthlySavings': {}, 'topMoversUp': [], 'topMoversDown': [],
            'realizedSavings': 0.0, 'exceptionFloor': 0.0,
            'nextSteps': [], 'ongoingNextSteps': [],
            'plannedSavings': [], 'projectUpdates': [], 'progressNarrative': '',
        }

    blob_path = blob_client.upload_docx(customer_id, month, year, file_bytes, filename)

    report_id   = str(uuid.uuid4())
    generated_at = datetime.now(timezone.utc)
    report = Report(
        id=report_id,
        customerId=customer_id,
        month=month,
        year=year,
        status='imported',
        blobPath=blob_path,
        generatedAt=generated_at,
        joelNotes=None,
        narrativeDraft=None,
        source='manual_import',
        extractedData=extracted_data,
    )
    cosmos_client.create_report(report)

    logging.info(
        'Imported report %s for %s/%d/%d — monthlySavings=%s nextSteps=%d '
        'plannedSavings=%d projectUpdates=%d ongoingNextSteps=%d',
        report_id, customer_id, month, year,
        list(extracted_data['monthlySavings'].keys()),
        len(extracted_data['nextSteps']),
        len(extracted_data['plannedSavings']),
        len(extracted_data['projectUpdates']),
        len(extracted_data['ongoingNextSteps']),
    )

    return _json({
        'success': True,
        'reportId': report_id,
        'extractedData': extracted_data,
    })
